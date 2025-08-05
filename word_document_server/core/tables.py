"""
Table-related operations for Word Document Server.
"""
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import RGBColor


def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    
    Args:
        cell: The cell to modify
        **kwargs: Border properties (top, bottom, left, right, val, color)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for key, value in kwargs.items():
        if key in ['top', 'left', 'bottom', 'right']:
            tag = 'w:{}'.format(key)
            
            element = OxmlElement(tag)
            element.set(qn('w:val'), kwargs.get('val', 'single'))
            element.set(qn('w:sz'), kwargs.get('sz', '4'))
            element.set(qn('w:space'), kwargs.get('space', '0'))
            element.set(qn('w:color'), kwargs.get('color', 'auto'))
            
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
                
            tcBorders.append(element)


def apply_table_style(table, has_header_row=False, border_style=None, shading=None):
    """
    Apply formatting to a table.
    
    Args:
        table: The table to format
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick')
        shading: 2D list of cell background colors (by row and column)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Format header row if requested
        if has_header_row and table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Apply border style if specified
        if border_style:
            val_map = {
                'none': 'nil',
                'single': 'single',
                'double': 'double',
                'thick': 'thick'
            }
            val = val_map.get(border_style.lower(), 'single')
            
            # Apply to all cells
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top=True,
                        bottom=True,
                        left=True,
                        right=True,
                        val=val,
                        color="000000"
                    )
        
        # Apply cell shading if specified
        if shading:
            for i, row_colors in enumerate(shading):
                if i >= len(table.rows):
                    break
                for j, color in enumerate(row_colors):
                    if j >= len(table.rows[i].cells):
                        break
                    try:
                        # Apply shading to cell
                        cell = table.rows[i].cells[j]
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        # Skip if color format is invalid
                        pass
        
        return True
    except Exception:
        return False


def copy_table(source_table, target_doc):
    """
    Copy a table from one document to another.
    
    Args:
        source_table: The table to copy
        target_doc: The document to copy the table to
        
    Returns:
        The new table in the target document
    """
    # Create a new table with the same dimensions
    new_table = target_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    
    # Try to apply the same style
    try:
        if source_table.style:
            new_table.style = source_table.style
    except:
        # Fall back to default grid style
        try:
            new_table.style = 'Table Grid'
        except:
            pass
    
    # Copy cell contents
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                if paragraph.text:
                    new_table.cell(i, j).text = paragraph.text
    
    return new_table


def set_cell_shading(cell, fill_color=None, pattern="clear", pattern_color="auto"):
    """
    Apply shading/filling to a table cell.
    
    Args:
        cell: The table cell to format
        fill_color: Background color (hex string like "FF0000" or RGBColor)
        pattern: Shading pattern ("clear", "solid", "pct10", "pct20", etc.)
        pattern_color: Pattern color for patterned fills
        
    Returns:
        True if successful, False otherwise
    """
    try:
        # Get or create table cell properties
        tc_pr = cell._tc.get_or_add_tcPr()
        
        # Remove existing shading
        existing_shd = tc_pr.find(qn('w:shd'))
        if existing_shd is not None:
            tc_pr.remove(existing_shd)
        
        # Create shading element
        shd_attrs = {
            'w:val': pattern,
            'w:color': pattern_color if pattern_color != "auto" else "auto"
        }
        
        # Set fill color
        if fill_color:
            if isinstance(fill_color, str):
                # Hex color string - remove # if present
                fill_color = fill_color.lstrip('#').upper()
                if len(fill_color) == 6:  # Valid hex color
                    shd_attrs['w:fill'] = fill_color
            elif isinstance(fill_color, RGBColor):
                # RGBColor object
                hex_color = f"{fill_color.r:02X}{fill_color.g:02X}{fill_color.b:02X}"
                shd_attrs['w:fill'] = hex_color
        
        # Build XML string
        attr_str = ' '.join([f'{k}="{v}"' for k, v in shd_attrs.items()])
        shd_xml = f'<w:shd {nsdecls("w")} {attr_str}/>'
        
        # Parse and append shading element
        shading_elm = parse_xml(shd_xml)
        tc_pr.append(shading_elm)
        
        return True
        
    except Exception as e:
        print(f"Error setting cell shading: {e}")
        return False


def apply_alternating_row_shading(table, color1="FFFFFF", color2="F2F2F2"):
    """
    Apply alternating row colors for better readability.
    
    Args:
        table: The table to format
        color1: Color for odd rows (hex string)
        color2: Color for even rows (hex string)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        for i, row in enumerate(table.rows):
            fill_color = color1 if i % 2 == 0 else color2
            for cell in row.cells:
                set_cell_shading(cell, fill_color=fill_color)
        return True
    except Exception as e:
        print(f"Error applying alternating row shading: {e}")
        return False


def highlight_header_row(table, header_color="4472C4", text_color="FFFFFF"):
    """
    Apply special shading to header row.
    
    Args:
        table: The table to format
        header_color: Background color for header (hex string)
        text_color: Text color for header (hex string)
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if table.rows:
            for cell in table.rows[0].cells:
                # Apply background shading
                set_cell_shading(cell, fill_color=header_color)
                
                # Apply text formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        if text_color and text_color != "auto":
                            # Convert hex to RGB
                            try:
                                text_color = text_color.lstrip('#')
                                r = int(text_color[0:2], 16)
                                g = int(text_color[2:4], 16)
                                b = int(text_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                            except:
                                pass  # Skip if color format is invalid
        return True
    except Exception as e:
        print(f"Error highlighting header row: {e}")
        return False


def set_cell_shading_by_position(table, row_index, col_index, fill_color, pattern="clear"):
    """
    Apply shading to a specific cell by row/column position.
    
    Args:
        table: The table containing the cell
        row_index: Row index (0-based)
        col_index: Column index (0-based)
        fill_color: Background color (hex string)
        pattern: Shading pattern
        
    Returns:
        True if successful, False otherwise
    """
    try:
        if (0 <= row_index < len(table.rows) and 
            0 <= col_index < len(table.rows[row_index].cells)):
            cell = table.rows[row_index].cells[col_index]
            return set_cell_shading(cell, fill_color=fill_color, pattern=pattern)
        else:
            return False
    except Exception as e:
        print(f"Error setting cell shading by position: {e}")
        return False
